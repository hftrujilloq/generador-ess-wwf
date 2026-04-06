from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from clauses_library import CLAUSULAS


class DocumentGenerator:
    def __init__(self, nombre_proyecto, objeto_contrato, territorio):
        self.nombre_proyecto = nombre_proyecto
        self.objeto_contrato = objeto_contrato
        self.territorio = territorio
        self.doc = None

    def _crear_documento_base(self):
        self.doc = Document()

    def generar_anexo(self, lista_clausulas, codigos_activados):
        """
        Genera el anexo SAS en Word.

        Parámetros:
            lista_clausulas (list[str]): lista de IDs de cláusulas, por ejemplo
                ["SAS-00", "SAS-01", "ME&A-01", "PI-01"]
            codigos_activados (list[str]): lista de activadores SAS, por ejemplo
                ["A-PI", "A-ED"]
        """
        self._crear_documento_base()

        # Título principal
        titulo = self.doc.add_heading("ANEXO SAS – CLÁUSULAS CONTRACTUALES APLICABLES", 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Información general
        self.doc.add_paragraph(f"Proyecto: {self.nombre_proyecto}")
        self.doc.add_paragraph(f"Objeto contractual: {self.objeto_contrato}")
        self.doc.add_paragraph(f"Territorio: {self.territorio}")
        self.doc.add_paragraph(
            f"Activadores SAS identificados: {', '.join(codigos_activados) if codigos_activados else 'Ninguno'}"
        )
        self.doc.add_paragraph()

        # Introducción ajustada a la matriz vigente
        intro = self.doc.add_paragraph()
        intro.add_run(
            "Este anexo forma parte integral del contrato/acuerdo/convenio y documenta las "
            "cláusulas contractuales aplicables derivadas del Cuestionario de Screening SAS "
            "y de la Matriz de Decisión SAS de WWF Colombia. "
        )
        intro.add_run(
            "De conformidad con dicha matriz, cada respuesta afirmativa activa un código SAS, "
            "y cada código activa automáticamente un paquete cerrado de cláusulas contractuales. "
            "Las cláusulas aquí incorporadas son de obligatorio cumplimiento."
        )

        self.doc.add_paragraph()

        # Caso sin activaciones
        if not lista_clausulas:
            p = self.doc.add_paragraph()
            p.add_run(
                "No se activaron cláusulas SAS con base en las respuestas suministradas "
                "en el cuestionario de screening."
            )
            p.paragraph_format.space_after = Pt(12)
        else:
            # Cláusulas
            for codigo in lista_clausulas:
                if codigo not in CLAUSULAS:
                    continue

                clausula = CLAUSULAS[codigo]

                # Título de cláusula
                heading = self.doc.add_heading(clausula["titulo"], level=2)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Texto de cláusula
                parrafo = self.doc.add_paragraph(clausula["texto"])
                parrafo.paragraph_format.space_after = Pt(12)
                parrafo.paragraph_format.first_line_indent = Inches(0.25)

        # Cierre
        self.doc.add_paragraph()
        cierre = self.doc.add_paragraph("— Fin del Anexo SAS —")
        cierre.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return self.doc

    def descargar_word(self):
        if self.doc is None:
            raise ValueError("Primero debe ejecutar generar_anexo().")

        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer

    def guardar_local(self, filename):
        if self.doc is None:
            raise ValueError("Primero debe ejecutar generar_anexo().")

        self.doc.save(filename)
